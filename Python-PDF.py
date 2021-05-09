import os
import re
import pyperclip
from docx import Document

os.chdir('c:\\Users\\User\\Desktop\\z')

#Copy Content
content = pyperclip.paste()

#Create Word Document
wordDoc = Document()

#extract questions
regex = re.compile(r'((\d)+(\.)( )*[a-zA-Z0-9-%’=():, ]+(\?)*)')
try:
    extractedQuestion = regex.findall(content)
    allQuestions = []
    for question in extractedQuestion:
        allQuestions.append(question[0])
    questions = '\n'.join(allQuestions)
except AttributeError:
    print("No match found")

#extract answers
from typing import List, Tuple

import fitz  # install with 'pip install pymupdf'
answers = []

def _parse_highlight(annot: fitz.Annot, wordlist: List[Tuple[float, float, float, float, str, int, int, int]]) -> str:
    points = annot.vertices
    quad_count = int(len(points) / 4)
    sentences = []
    for i in range(quad_count):
        # where the highlighted part is
        r = fitz.Quad(points[i * 4 : i * 4 + 4]).rect

        words = [w for w in wordlist if fitz.Rect(w[:4]).intersects(r)]
        sentences.append(" ".join(w[4] for w in words))
    sentence = " ".join(sentences)
    return sentence


def handle_page(page):
    wordlist = page.getText("words")  # list of words on page
    wordlist.sort(key=lambda w: (w[3], w[0]))  # ascending y, then x

    highlights = []
    annot = page.firstAnnot
    while annot:
        if annot.type[0] == 8:
            highlights.append(_parse_highlight(annot, wordlist))
        annot = annot.next
    return highlights


doc = fitz.open('c:\\Users\\User\\Desktop\\z\\8-Virology.pdf')

highlights = []
for page in doc:
    highlights += handle_page(page)

#1st answer is at 3rd answer, so change position

allAnswers = highlights
ans = '\n\n'.join(highlights)


for i in range(0,len(allQuestions)):
    p = wordDoc.add_heading(allQuestions[i], level=2)
    p = wordDoc.add_paragraph(allAnswers[i])
    
wordDoc.save('8-Virology.docx')
print("Doc saved")
